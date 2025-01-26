import os
from dotenv import load_dotenv
from datetime import datetime
import httpx
from deepgram import DeepgramClient, DeepgramClientOptions, PrerecordedOptions, FileSource
from docx import Document  # For Word document creation
import streamlit as st
from tempfile import NamedTemporaryFile

# Load environment variables from .env file
load_dotenv()

def save_to_word_file(transcription_text: str, output_filename: str):
    """
    Save the transcription text to a Word document.
    """
    document = Document()
    document.add_heading("Transcription", level=1)  # Add a heading
    document.add_paragraph(transcription_text)  # Add the transcription text
    document.save(output_filename)
    return output_filename

def filter_duplicates(utterances):
    """
    Filter out duplicate or overlapping utterances based on timestamps and text similarity.
    """
    filtered = []
    seen_texts = set()

    for utterance in utterances:
        text = utterance.transcript.strip()
        start_time = utterance.start
        end_time = utterance.end

        # Skip if text is already seen or overlaps significantly with the last added utterance
        if text in seen_texts:
            continue
        if filtered and abs(start_time - filtered[-1]["end_time"]) < 0.5:
            continue

        # Add to filtered list
        filtered.append({
            "start_time": start_time,
            "end_time": end_time,
            "speaker": getattr(utterance, "speaker", "Unknown Speaker"),
            "text": text
        })
        seen_texts.add(text)

    return filtered

def format_diarized_transcription(utterances=None, words=None):
    """
    Format the transcription data without timestamps, grouping consecutive utterances by the same speaker.
    """
    speaker_transcriptions = []
    current_speaker = None
    current_text = []

    if utterances:
        filtered_utterances = filter_duplicates(utterances)
        for utterance in filtered_utterances:
            speaker = utterance["speaker"]
            text = utterance["text"]

            if speaker == current_speaker:
                # Append to the current speaker's text block
                current_text.append(text)
            else:
                # Save the previous speaker's text block
                if current_speaker is not None:
                    speaker_transcriptions.append(
                        f"{current_speaker}: {' '.join(current_text)}"
                    )
                # Start a new block for the new speaker
                current_speaker = speaker
                current_text = [text]

        # Add the last speaker's block
        if current_speaker is not None:
            speaker_transcriptions.append(
                f"{current_speaker}: {' '.join(current_text)}"
            )
    elif words:  # Fallback to word-level data if no utterances
        seen_texts = set()
        for word in words:
            text = word.punctuated_word.strip()
            if text not in seen_texts:
                speaker_transcriptions.append(
                    f"Unknown Speaker: {text}"
                )
                seen_texts.add(text)

    return "\n\n".join(speaker_transcriptions)

def main():
    st.title("Transcendental Transcriber")

    # Upload file widget
    uploaded_file = st.file_uploader("Upload your audio file", type=["wav", "mp4", "mp3"])

    if uploaded_file:
        st.success("File uploaded successfully!")

        # Load Deepgram API key
        api_key = os.getenv("DEEPGRAM_API_KEY")
        if not api_key:
            st.error("Deepgram API key not found. Please set it in the .env file.")
            return

        try:
            # Create a temporary file to save the uploaded file
            with NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(uploaded_file.read())
                temp_file_path = temp_file.name

            # Create Deepgram client
            config: DeepgramClientOptions = DeepgramClientOptions(verbose=False)
            deepgram: DeepgramClient = DeepgramClient(api_key, config)

            # Read the audio file
            with open(temp_file_path, "rb") as file:
                buffer_data = file.read()

            # Prepare payload and options
            payload: FileSource = {"buffer": buffer_data}
            options: PrerecordedOptions = PrerecordedOptions(
                model="general",
                smart_format=True,
                utterances=True,
                punctuate=True,
                diarize=True,
                multichannel=True
            )

            # Transcribe the audio file
            response = deepgram.listen.rest.v("1").transcribe_file(
                payload, options, timeout=httpx.Timeout(300.0, connect=10.0)
            )

            # Access the transcription results
            if hasattr(response, "results") and response.results.channels:
                channel = response.results.channels[0]
                alternatives = channel.alternatives

                if alternatives:
                    transcript = alternatives[0].transcript
                    utterances = getattr(response.results, "utterances", [])
                    words = getattr(alternatives[0], "words", [])

                    # Use utterances if available; otherwise, use words
                    if utterances:
                        formatted_transcription = format_diarized_transcription(utterances=utterances)
                    elif words:
                        formatted_transcription = format_diarized_transcription(words=words)
                    else:
                        formatted_transcription = transcript

                    # Save to a Word file
                    output_filename = f"{uploaded_file.name.split('.')[0]}_transcription.docx"
                    save_path = save_to_word_file(formatted_transcription, output_filename)

                    # Provide download link
                    st.success("Transcription completed!")
                    st.download_button(
                        label="Download Transcription",
                        data=open(save_path, "rb").read(),
                        file_name=output_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("No transcription alternatives found.")
            else:
                st.error("No results found in the transcription response.")

        except Exception as e:
            st.error(f"Error: {e}")

if __name__ == "__main__":
    main()
